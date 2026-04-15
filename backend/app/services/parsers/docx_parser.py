from docx import Document
import io

from app.config import settings


class DocxParser:
    """
    Extract text from Word (.docx) files via python-docx.
    Tables are rendered as pipe-delimited rows.
    Word has no real "pages" — we simulate them at ~3000 chars.
    """

    def extract(self, file_bytes: bytes) -> list[dict]:
        doc = Document(io.BytesIO(file_bytes))
        pages = []
        current_text = ""
        char_count = 0
        page_char_limit = max(500, settings.extract_virtual_page_char_limit)

        for para in doc.paragraphs:
            current_text += para.text + "\n"
            char_count += len(para.text)
            if char_count >= page_char_limit:
                pages.append({"page": len(pages) + 1, "text": current_text})
                current_text = ""
                char_count = 0

        if current_text.strip():
            pages.append({"page": len(pages) + 1, "text": current_text})

        for table in doc.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            table_text = "\n".join(rows)
            if table_text.strip():
                pages.append({"page": len(pages) + 1, "text": f"[Table]\n{table_text}"})

        return pages
